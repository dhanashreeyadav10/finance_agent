# from docx import Document

# def generate_doc_report(insights):
#     doc = Document()
#     doc.add_heading("Finance AI Report", level=1)

#     for k, v in insights.items():
#         doc.add_heading(k, level=2)
#         doc.add_paragraph(v)

#     path = "finance_report.docx"
#     doc.save(path)
#     return path

from docx import Document

def generate_doc_report(insights):
    doc = Document()
    doc.add_heading("Finance AI Report", level=1)

    for k, v in insights.items():
        doc.add_heading(k, level=2)
        doc.add_paragraph(v)

    path = "Finance_Report.docx"
    doc.save(path)
    return path
